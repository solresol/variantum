from __future__ import annotations

import csv
import math
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/gregb/Documents/devel/variantum")
OUT = ROOT / "outputs" / "ai4as-2026-parallage"
DATA = ROOT / "analysis" / "vanessa-set1-metric-scatter-data.csv"
RAW_PLOT = ROOT / "analysis" / "vanessa-set1-raw-metric-scatter.png"
ADJUSTED_PLOT = ROOT / "analysis" / "vanessa-set1-length-adjusted-residual-scatter.png"
LENGTH_PLOT = ROOT / "analysis" / "vanessa-set1-length-and-composite-scatter.png"
POP_PLOT = ROOT / "analysis" / "vanessa-set1-population-length-bias-scatter.png"
DOCX = OUT / "parallage_ai4as_2026_speaker_brief.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "E8EEF5"


def load_rows() -> list[dict[str, str]]:
    with DATA.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def f(value: str) -> float:
    return float(value)


def spearman(xs: list[float], ys: list[float]) -> float:
    def ranks(values: list[float]) -> list[float]:
        order = sorted(range(len(values)), key=lambda i: values[i])
        out = [0.0] * len(values)
        i = 0
        while i < len(order):
            j = i
            while j + 1 < len(order) and values[order[j + 1]] == values[order[i]]:
                j += 1
            rank = (i + j + 2) / 2
            for k in range(i, j + 1):
                out[order[k]] = rank
            i = j + 1
        return out

    rx = ranks(xs)
    ry = ranks(ys)
    mx = sum(rx) / len(rx)
    my = sum(ry) / len(ry)
    num = sum((x - mx) * (y - my) for x, y in zip(rx, ry))
    den = math.sqrt(sum((x - mx) ** 2 for x in rx) * sum((y - my) ** 2 for y in ry))
    return num / den if den else float("nan")


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, *, size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc: Document, text: str = "", *, style=None, bold=False, italic=False, color=None, size=None, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.4)
    set_cell_fill(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=DARK_BLUE, size=11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5)
    doc.add_paragraph()


def add_small_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for idx, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        cell.width = Inches(width)
        set_cell_fill(cell, LIGHT_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        set_run_font(r, bold=True, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, (cell, text, width) in enumerate(zip(cells, row, widths)):
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=9.2)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(5.85))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(caption)
        set_run_font(r, size=9.5, color=MUTED, italic=True)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "AI4AS 2026 speaker brief"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "Draft based on published abstract, Parallage project docs, and Vanessa Set 1 pilot analysis"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED


def build() -> None:
    rows = load_rows()
    ratings = [f(r["rating"]) for r in rows]
    raw_bleu = [f(r["mean_bleu4"]) for r in rows]
    raw_rouge = [f(r["mean_rouge_l"]) for r in rows]
    raw_trigram = [f(r["mean_3gram_f1"]) for r in rows]
    resid_bleu = [f(r["mean_bleu4_resid_badness"]) for r in rows]
    resid_rouge = [f(r["mean_rouge_l_resid_badness"]) for r in rows]
    resid_trigram = [f(r["mean_3gram_f1_resid_badness"]) for r in rows]
    lengths = [f(r["reference_words"]) for r in rows]

    high = [r for r in rows if f(r["rating"]) >= 7]
    top_resid = sorted(rows, key=lambda r: f(r["composite_resid_badness"]), reverse=True)[:4]

    doc = Document()
    configure_styles(doc)

    add_para(doc, "AI4AS 2026 Parallage Talk", bold=True, size=23, color=RGBColor(0, 0, 0), after=4)
    add_para(
        doc,
        "What to say and show for Into the Parallage: Harnessing Abundance, Plurality and Divergence in AI Translation of Ancient Texts",
        size=13,
        color=MUTED,
        after=14,
    )
    add_small_table(
        doc,
        ["Field", "Current evidence"],
        [
            ["Venue", "AI4AS 2026: Artificial Intelligence, Engagement, and the Fragility of Ancient Scripts, online mini-conference within DH2026."],
            ["Date", "27 July 2026; scheduled in the 11:00-12:00 session on AI and Translations."],
            ["Authors", "Greg Baker, Shirley Chan, Vanessa Enriquez Raido, Greta Hawes."],
            ["Talk promise", "The abstract commits us to theoretical framework, corpus rationale, and preliminary pilot findings."],
            ["Main pilot evidence", "Vanessa Set 1: 10 current ratings against hidden human translations for Stephanos v3 outputs, with sentence-aligned lexical metrics and length-adjusted residuals."],
        ],
        [1.35, 5.0],
    )

    add_callout(
        doc,
        "Recommended claim",
        "Parallage is not an attempt to produce one better AI translation. It is an interface and method for making translation multiplicity inspectable, so that disagreement and uncertainty become evidence rather than noise.",
    )

    doc.add_heading("1. The talk we should give", level=1)
    add_para(
        doc,
        "The published abstract points to a talk that is less about proving a finished evaluation result and more about showing a defensible research object: structured translation packs for low-resource ancient texts, with a first pilot that tests whether human judgement and automatic overlap signals are looking at the same fragility.",
    )
    add_bullets(
        doc,
        [
            "Open with the problem: single fluent translations hide the plural interpretive work that produced them.",
            "Introduce Parallage as a deliberate use of AI abundance: multiple situated renderings plus inspection cues.",
            "Ground the method in two corpora: Stephanos of Byzantium's Ethnica and Hanshu Dilizhi.",
            "Use Vanessa's pilot as preliminary evidence, not overclaiming: she identifies some high-divergence cases, but raw overlap metrics are length-biased.",
            "Close with the next study: more reviewers, Chinese material, and direct comparison between single-output and pack-based interfaces.",
        ]
    )

    doc.add_heading("2. Ten-minute talk script", level=1)
    script_sections = [
        (
            "Opening",
            "Translation is often presented to readers as a finished object. That presentation is misleading. Expert translators do not retrieve a single correct answer; they move through alternatives, constraints, purposes, and risks. The final translation is only the visible residue of a plural process.",
        ),
        (
            "Why AI changes the problem",
            "AI makes translation abundant, but abundance alone does not make translation more reliable. A fluent model output can compress uncertainty into a single authoritative surface. In ancient-language work, especially where texts are fragmentary, culturally dense, or under-translated, that is exactly the wrong direction.",
        ),
        (
            "Parallage",
            "Parallage asks whether we can structure AI abundance so that it exposes rather than conceals interpretive choice. Instead of one AI translation, a reader receives a pack: literal, readable, interpretive, uncertainty-marked, and adversarial renderings, with lightweight cues for source segmentation, named entities, and ambiguity.",
        ),
        (
            "Corpus rationale",
            "The case studies are Stephanos of Byzantium's Ethnica and Hanshu Dilizhi. Neither has a complete English translation in the target form used here. That matters methodologically because readers and semi-expert translators cannot simply defer to a stable English canon; trust and interpretive agency become visible tasks.",
        ),
        (
            "Pilot",
            "The preliminary pilot uses the Greek side. Vanessa reviewed the first set of ten Stephanos passages and rated how different she expected the hidden human translation to be from the focal automated translation. We then compared those ratings with sentence-aligned BLEU, ROUGE-L, and n-gram overlap scores against the approved human translations.",
        ),
        (
            "Raw result",
            f"At first pass, there is a directional signal: higher ratings tend to coincide with lower overlap scores. The raw Spearman correlations are about {spearman(ratings, raw_bleu):.2f} for BLEU, {spearman(ratings, raw_rouge):.2f} for ROUGE-L, and {spearman(ratings, raw_trigram):.2f} for 3-gram F1 when correlating rating against the raw score. The sign is expected to be negative.",
        ),
        (
            "Length caveat",
            "But these metrics punish longer passages. Across the broader v3 metric population, longer reference translations score worse, and Vanessa's ratings in Set 1 also rise with length. So the raw signal is partly a length signal.",
        ),
        (
            "Adjusted result",
            f"After adjusting expected metric score by reference length, the residual signal is weaker: about {spearman(ratings, resid_bleu):.2f} for BLEU residual badness, {spearman(ratings, resid_rouge):.2f} for ROUGE-L residual badness, and {spearman(ratings, resid_trigram):.2f} for 3-gram F1 residual badness. This is suggestive, not conclusive.",
        ),
        (
            "Interpretation",
            "The useful result is not that Vanessa perfectly predicted automatic metrics. She did not. The useful result is that the review task produces analyzable judgement data, and that the mismatch between reviewer judgement, metric scores, and passage length tells us where the next study must be more careful.",
        ),
        (
            "Close",
            "The current evidence supports Parallage as a publishable method-in-progress: a way to transform AI-generated multiplicity into a structured object for judgement. The next step is to scale from one expert reviewer and ten Greek passages to multiple reviewers, the Chinese stream, and experimental comparison against single-output translation.",
        ),
    ]
    for title, body in script_sections:
        doc.add_heading(title, level=2)
        add_para(doc, body)

    doc.add_heading("3. Pilot results to say aloud", level=1)
    add_small_table(
        doc,
        ["Measure", "Raw signal", "Length-adjusted signal", "How to phrase it"],
        [
            ["BLEU", f"rating vs score r={spearman(ratings, raw_bleu):.2f}", f"rating vs residual badness r={spearman(ratings, resid_bleu):.2f}", "Weak after length control."],
            ["ROUGE-L", f"rating vs score r={spearman(ratings, raw_rouge):.2f}", f"rating vs residual badness r={spearman(ratings, resid_rouge):.2f}", "Raw signal mostly does not survive strongly."],
            ["3-gram F1", f"rating vs score r={spearman(ratings, raw_trigram):.2f}", f"rating vs residual badness r={spearman(ratings, resid_trigram):.2f}", "Most suggestive adjusted metric, still n=10."],
            ["Length", f"ratings rise with length; reference words range {int(min(lengths))}-{int(max(lengths))}", "Length model trained on 101 v3 rows.", "Say this as a caveat, not a footnote."],
        ],
        [1.1, 1.45, 1.7, 2.05],
    )

    add_para(doc, "High-difference cases Vanessa selected:", bold=True, color=DARK_BLUE)
    add_bullets(doc, [f"{r['tier_rank']}: {r['lemma_display']} (rating {r['rating']})" for r in high])
    add_para(doc, "Largest length-adjusted residual badness cases:", bold=True, color=DARK_BLUE)
    add_bullets(
        doc,
        [
            f"{r['tier_rank']}: {r['lemma_display']} (rating {r['rating']}, composite residual {float(r['composite_resid_badness']):.2f})"
            for r in top_resid
        ],
    )

    doc.add_heading("4. Figures for deck or paper", level=1)
    add_image(doc, RAW_PLOT, "Figure 1. Raw metric scores against Vanessa rating. Higher ratings generally mean expected divergence; lower scores mean lower overlap.")
    add_image(doc, ADJUSTED_PLOT, "Figure 2. Length-adjusted residual badness. Positive values mean a translation scored worse than expected for reference length.")
    add_image(doc, LENGTH_PLOT, "Figure 3. Ratings track length, and composite adjusted badness weakens the raw story.")
    add_image(doc, POP_PLOT, "Figure 4. Broader v3 metric population shows the length bias that must be controlled.")

    doc.add_heading("5. Claims to avoid", level=1)
    add_bullets(
        doc,
        [
            "Do not claim that the pilot proves Parallage improves reliability; there is no control condition yet.",
            "Do not claim that Vanessa reliably found the lowest BLEU/ROUGE/n-gram translations; the length-adjusted evidence is weak at n=10.",
            "Do not treat overlap metrics as ground truth for translation quality; they are diagnostic proxies against an approved human reference.",
            "Do not claim the Chinese side has been evaluated unless the Hanshu material is actually ready before the talk.",
        ]
    )

    doc.add_heading("6. Source basis", level=1)
    add_bullets(
        doc,
        [
            "AI4AS 2026 conference page and programme: online mini-conference within DH2026, 27 July 2026.",
            "Published abstract: Baker, Chan, Enriquez Raido, Hawes, 'Into the Parallage: Harnessing Abundance, Plurality and Divergence in AI Translation of Ancient Texts'.",
            "Variantum README and Parallage proposal plan for research questions, roles, corpus rationale, and pilot structure.",
            "Local Vanessa Set 1 analysis in variantum/analysis, generated from the live Parallage review DB and Stephanos sentence metrics.",
        ]
    )

    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
